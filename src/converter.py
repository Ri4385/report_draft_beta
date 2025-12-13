import tempfile
from pathlib import Path
import re

import pypandoc  # type: ignore


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    """
    マークダウンをdocsに変換する。
    mdのpathの親にmediaがある想定。

    ### 処理内容
    1. mdの前処理
    2. pandocでmd→docx変換
    3. docxの後処理

    ### 処理内容詳細
    - 文中の$$ x $$ → $ x $
    - 行全体が$$ x $$の前後に空行追加
    - 表の前後に空行追加
    - 図の前後に空行追加
    - docsに変換し保存
    - docxのフォントカラーを黒に変更
    - docxのフォントを明朝に変更
    - docsで段落を両端揃えに変更
    - docxで表を中央揃えに変更
    - 表の直前の段落（キャプション）を中央揃えに変更
    - docxの表罫線を上下に追加
    - docs上書き保存
    """
    # TODO: 関数分離するなどコード整理をする.
    text = md_path.read_text(encoding="utf-8")

    lines = text.splitlines()

    # 方針
    # 1) マルチライン・ディスプレイ数式ブロック（$$ のみの行で囲う）は保持し、前後に空行を付与
    # 2) 1行完結のディスプレイ数式（行全体が $$ ... $$ の場合）も保持し、前後に空行を付与
    # 3) 文中に埋め込まれた $$...$$ は $...$ に変換（インライン化）
    #    → 具体的には、行頭・行末以外に文字がある（=文中）場合のみ $...$ に置換
    #
    # 4) 表は既に良好とのことなので、表の空行整形は維持（必要最小限）

    # まず、マルチライン・ディスプレイブロックの範囲を検出
    display_block_lines: set[int] = set()
    i = 0
    while i < len(lines):
        if re.match(r"^\s*\$\$\s*$", lines[i]):
            start = i
            i += 1
            while i < len(lines) and not re.match(r"^\s*\$\$\s*$", lines[i]):
                i += 1
            if i < len(lines):  # closing $$ found
                end = i
                for k in range(start, end + 1):
                    display_block_lines.add(k)
            # i は最後にインクリメントされる
        i += 1

    processed: list[str] = []
    j = 0
    while j < len(lines):
        line = lines[j]

        # マルチライン・ディスプレイブロックはそのまま移送し、前後に空行を付与
        if j in display_block_lines:
            # ブロック先頭であれば、直前に空行を追加
            if len(processed) > 0 and processed[-1].strip() != "":
                processed.append("")
            # ブロック全体をコピー
            while j < len(lines) and j in display_block_lines:
                processed.append(lines[j])
                j += 1
            # ブロック終端の直後に空行
            if j < len(lines) and lines[j].strip() != "":
                processed.append("")
            continue

        # 1行完結のディスプレイ数式（行全体が $$...$$）を検出して保持
        m_line_display = re.match(r"^\s*\$\$(.+?)\$\$\s*$", line)
        if m_line_display:
            if len(processed) > 0 and processed[-1].strip() != "":
                processed.append("")
            processed.append(line.strip())  # そのまま
            # 次行に空行が無ければ入れる
            if j + 1 < len(lines) and lines[j + 1].strip() != "":
                processed.append("")
            j += 1
            continue

        # 文中の $$...$$ → $...$ に置換（文中判定）
        # 文中判定：前後に非空白文字が存在する $$...$$
        # まず、すべての $$...$$ を候補抽出
        def inline_replacer(match: re.Match[str]) -> str:
            inner = match.group(1)
            return f"${inner}$"

        # 置換は行中に他のテキストがある場合のみ行う
        if "$$" in line:
            # 1) まず、文中の $$...$$ に限定して置換（貪欲回避）
            #    ただし、行頭直後から始まり行末直前で終わる（=行全体が $$...$$）ケースは上で捕捉済み
            newline = re.sub(r"\$\$(.+?)\$\$", inline_replacer, line)
            processed.append(newline)
        else:
            processed.append(line)
        j += 1

    # 表まわりの空行整形（最小限）：表ブロックの前後に空行
    final_lines: list[str] = []
    k = 0
    while k < len(processed):
        line = processed[k]
        if line.lstrip().startswith("|"):
            if len(final_lines) > 0 and final_lines[-1].strip() != "":
                final_lines.append("")
            while k < len(processed) and processed[k].lstrip().startswith("|"):
                final_lines.append(processed[k])
                k += 1
            if k < len(processed) and processed[k].strip() != "":
                final_lines.append("")
            continue
        if line.lstrip().startswith("![]("):
            if len(final_lines) > 0 and final_lines[-1].strip() != "":
                final_lines.append("")
            final_lines.append(line)
            if k + 1 < len(processed) and processed[k + 1].strip() != "":
                final_lines.append("")
            k += 1
            continue
        final_lines.append(line)
        k += 1

    final_md = "\n".join(final_lines)

    final_md = "\n".join(final_lines)

    # 一時ファイルに保存（方法1: with文を使用）
    with tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8", suffix=".md", delete=False
    ) as tmp_file:
        tmp_file.write(final_md)
        tmp_file_path = tmp_file.name

    # 変換（Pandoc）
    args_from = "markdown+tex_math_dollars+pipe_tables+grid_tables+table_captions"

    pypandoc.convert_file(  # type: ignore
        tmp_file_path,
        "docx",
        format=args_from,
        outputfile=docx_path,
        extra_args=[
            "--to=docx",
            "--mathml",
            "--standalone",
            "--columns=120",
            "--resource-path",
            # str(md_path.parent) + "/media",
            str(md_path.parent),
        ],
    )

    # 一時ファイルを削除
    Path(tmp_file_path).unlink()

    # TODO: ここから下のコード整理
    # あとどうのようなスタイル調整をするかも決める

    # docxのフォントカラーと表罫線を修正

    from docx import Document
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(docx_path))

    def set_run_black(run):  # type: ignore
        run.font.color.rgb = RGBColor(0, 0, 0)  # type: ignore
        # フォントを明朝に設定
        run.font.name = "MS Mincho"  # type: ignore
        r = run._element  # type: ignore
        rPr = r.get_or_add_rPr()  # type: ignore
        rFonts = rPr.find(qn("w:rFonts"))  # type: ignore
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")  # type: ignore
            rPr.append(rFonts)  # type: ignore
        rFonts.set(qn("w:eastAsia"), "MS Mincho")  # type: ignore 日本語用
        rFonts.set(qn("w:ascii"), "Times New Roman")  # type: ignore 英数字用
        rFonts.set(qn("w:hAnsi"), "Times New Roman")  # type: ignore

    for p in doc.paragraphs:
        # 段落を両端揃えに設定
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            set_run_black(r)

    def set_tbl_borders(table, top=True, bottom=True):  # type: ignore
        tbl = table._element  # type: ignore
        tblPr = tbl.tblPr  # type: ignore
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")  # type: ignore
            tbl.append(tblPr)  # type: ignore
        # Create or get tblBorders
        tblBorders = tblPr.find(qn("w:tblBorders"))  # type: ignore
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")  # type: ignore
            tblPr.append(tblBorders)  # type: ignore

        def ensure(tag, sz="8", val="single", color="000000", space="0"):  # type: ignore
            el = tblBorders.find(qn(f"w:{tag}"))  # type: ignore
            if el is None:
                el = OxmlElement(f"w:{tag}")  # type: ignore
                tblBorders.append(el)  # type: ignore
            el.set(qn("w:val"), val)  # type: ignore
            el.set(qn("w:sz"), sz)  # type: ignore
            el.set(qn("w:color"), color)  # type: ignore
            el.set(qn("w:space"), space)  # type: ignore

        if top:
            ensure("top")
        if bottom:
            ensure("bottom")

    def set_cell_bottom_border(cell, sz="8", val="single", color="000000", space="0"):  # type: ignore
        tcPr = cell._tc.get_or_add_tcPr()  # type: ignore
        tcBorders = tcPr.find(qn("w:tcBorders"))  # type: ignore
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")  # type: ignore
            tcPr.append(tcBorders)  # type: ignore
        bottom = tcBorders.find(qn("w:bottom"))  # type: ignore
        if bottom is None:
            bottom = OxmlElement("w:bottom")  # type: ignore
            tcBorders.append(bottom)  # type: ignore
        bottom.set(qn("w:val"), val)  # type: ignore
        bottom.set(qn("w:sz"), sz)  # type: ignore
        bottom.set(qn("w:color"), color)  # type: ignore
        bottom.set(qn("w:space"), space)  # type: ignore

    for i, table in enumerate(doc.tables):
        # 表全体を中央揃えに
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # # 表の直前の段落（キャプション）を中央揃えに
        table_element = table._element  # type: ignore
        prev_element = table_element.getprevious()  # type: ignore
        if prev_element is not None and prev_element.tag.endswith("}p"):  # type: ignore
            # 段落要素を見つけたら中央揃えに
            for p in doc.paragraphs:
                if p._element == prev_element:  # type: ignore
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break

        set_tbl_borders(table, top=True, bottom=True)
        if table.rows:
            for cell in table.rows[0].cells:
                set_cell_bottom_border(cell, sz="8")

    for section in doc.sections:
        for hdrftr in (section.header, section.footer):
            for p in hdrftr.paragraphs:
                for r in p.runs:
                    set_run_black(r)

    doc.save(str(docx_path))

    print(f"Converted {md_path} to {docx_path}")
    return


if __name__ == "__main__":
    md_path = Path("exp4_20251205/report.md")
    # docx_path = Path("others/kankyo_anzen.docx")
    file_name = md_path.stem + ".docx"
    docx_path = md_path.parent / file_name
    print(f"Converting {md_path} to {docx_path}...")
    convert_md_to_docx(md_path, docx_path)
